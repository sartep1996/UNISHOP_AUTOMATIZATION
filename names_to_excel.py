import docx
import pandas as pd
import openpyxl



def extract_company_names(wordfile):
    doc = docx.Document(wordfile)
    company_names = []

    current_company_name = ''
    
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if run.bold:
                # Extend the current company name by adding the bold run's text
                current_company_name += ' ' + run.text if current_company_name else run.text
            else:
                # If the run is not bold, it's a separate word, so add it to the list if the current company name is not empty
                if current_company_name:
                    company_names.append(current_company_name)
                    current_company_name = ''  # Reset the current company name

    if current_company_name:
        company_names.append(current_company_name)

    return company_names
# def extract_company_names(wordfile):
#     doc = docx.Document(wordfile)
#     company_names = []

    
#     current_company_name = ''
#     for paragraph in doc.paragraphs:
#         for run in paragraph.runs:
#             if run.bold:
#                 # Extend the current company name by adding the bold run's text
#                 current_company_name += ' ' + run.text if current_company_name else run.text
#             else:
#                 # If the run is not bold, it's a separate word, so add it to the list if the current company name is not empty
#                 if current_company_name:
#                     company_names.append(current_company_name)
#                     current_company_name = ''  # Reset the current company name

#     if current_company_name:
        
#         company_names.append(current_company_name)

#     return company_names

input_file = "C:\\Users\\PetrasAnksaitis\\Downloads\\PythonFiles\\COMPANYNAMES WORD.docx"


company_names_extracted = extract_company_names(input_file)

company_names = list(set(company_names_extracted))
print(len(company_names))

# for name in company_names_extracted:
#     print(name)

additional_clients_df = pd.DataFrame({'Words': company_names})
excel_file = "C:\\Users\\PetrasAnksaitis\\Downloads\\PythonFiles\testclient.xlsx"
start_row = 1
start_col = 1

additional_clients_df.to_excel(excel_file, index=False, startrow=start_row, startcol=start_col, )

